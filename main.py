import os
import json
import hashlib
from typing import List, Optional, Dict, Any
from datetime import datetime
from fastapi import HTTPException
from langchain.chains import LLMChain
from fastapi.responses import PlainTextResponse
from fastapi.responses import JSONResponse
from fastapi import FastAPI, File, UploadFile, Form
from pydantic import BaseModel
from dotenv import load_dotenv
import PyPDF2
import numpy as np
# модули проекта
from document_manager import DocumentManager
from document_processor import process_document_folder
from vector_storage import create_faiss_index, load_faiss_index, save_faiss_index, add_chunks_to_index

# Импорт компонентов LangChain и OpenAI
from langchain_openai import ChatOpenAI
from langchain_community.vectorstores import FAISS as LC_FAISS
from langchain.docstore.document import Document as LCDocument
from langchain.docstore.document import Document
from langchain.chains.conversational_retrieval.base import ConversationalRetrievalChain
from langchain.prompts.prompt import PromptTemplate
from langchain.embeddings.base import Embeddings
import warnings
warnings.filterwarnings("ignore", message="CropBox missing from /Page, defaulting to MediaBox")

from sentence_transformers import SentenceTransformer

# ЗАГРУЗКА ENV
load_dotenv()

OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
# Для учителей
DATA_FOLDER_TEACHERS = os.getenv("DATA_FOLDER", "data")
INDEXES_FOLDER_TEACHERS = os.getenv("INDEXES_FOLDER", "indexes")
# Для студентов
DATA_FOLDER_STUDENTS = os.getenv("DATA_FOLDER_STUD", "data_stud")
INDEXES_FOLDER_STUDENTS = os.getenv("INDEXES_FOLDER_STUD", "indexes_stud")

# ============================================================
# Класс эмбеддингов (на базе sentence_transformers)
# ============================================================
class MyEmbeddings(Embeddings):
    def __init__(self):
        from sentence_transformers import SentenceTransformer
        self.model = SentenceTransformer('all-MiniLM-L12-v2')
    def embed_documents(self, texts):
        return self.model.encode(texts, convert_to_numpy=True).tolist()
    def embed_query(self, text):
        return self.model.encode([text], convert_to_numpy=True)[0].tolist()

embeddings = MyEmbeddings()



class ChatHistoryResponse(BaseModel):
    session_id: str
    # now each item has an int id plus role/content
    history: List[Dict[str, Any]]

class ChatDeleteResponse(BaseModel):
    message: str


# ============================================================
# Функция загрузки/перестроения векторного хранилища
# (упрощённая версия с контрольной меткой (fingerprint))
# ============================================================
def load_or_rebuild_vectorstore(data_folder: str, indexes_folder: str, call_id: str = "") -> LC_FAISS:
    # Create folders if they don't exist
    os.makedirs(indexes_folder, exist_ok=True)
    fingerprint_file = os.path.join(indexes_folder, "index_fingerprint.json")
    index_path = os.path.join(indexes_folder, "index.faiss")
    metadata_path = os.path.join(indexes_folder, "document_metadata.json")

    # Calculate current fingerprint
    current_fingerprint = {}
    for root, _, files in os.walk(data_folder):
        for file in files:
            if file.lower().endswith(('.docx', '.pdf', '.txt')):
                path = os.path.join(root, file)
                current_fingerprint[path] = os.path.getmtime(path)
                #print(f"[DEBUG] Fingerprint includes: {path}")
    fingerprint_hash = hashlib.md5(json.dumps(current_fingerprint, sort_keys=True).encode()).hexdigest()

    # Read previous fingerprint
    previous_fingerprint = None
    if os.path.exists(fingerprint_file):
        try:
            with open(fingerprint_file, 'r') as f:
                previous_fingerprint = json.load(f)
        except (json.JSONDecodeError, IOError) as e:
            print(f"[ERROR] Could not read fingerprint file: {e}")

    print(f"[DEBUG] Current fingerprint: {fingerprint_hash}")
    print(f"[DEBUG] Previous fingerprint: {previous_fingerprint}")

    # Check if we can load existing index
    if (os.path.exists(index_path) and os.path.exists(metadata_path) and previous_fingerprint == fingerprint_hash):
        print(f"[DEBUG] Index file exists: {os.path.exists(index_path)}")
        print(f"[DEBUG] Metadata file exists: {os.path.exists(metadata_path)}")

        try:
            # Here's the fix - wrap this in a more specific try/except block
            # that catches only errors we expect from FAISS loading
            try:
                vectorstore = LC_FAISS.load_local(indexes_folder, embeddings, allow_dangerous_deserialization=True)
                print(f"[INFO] Loaded existing vectorstore with {vectorstore.index.ntotal} vectors")
                return vectorstore
            except (AttributeError, ValueError, FileNotFoundError) as e:
                print(f"[ERROR] Failed to load existing vectorstore (specific error): {e}")
                # Delete corrupted files
                if os.path.exists(index_path):
                    print(f"[DEBUG] Removing corrupted index: {index_path}")
                    os.remove(index_path)
                if os.path.exists(metadata_path):
                    print(f"[DEBUG] Removing corrupted metadata: {metadata_path}")
                    os.remove(metadata_path)
        except Exception as e:
            print(f"[ERROR] Unexpected error loading vectorstore: {e}")
            # Delete corrupted files
            if os.path.exists(index_path):
                print(f"[DEBUG] Removing corrupted index: {index_path}")
                os.remove(index_path)
            if os.path.exists(metadata_path):
                print(f"[DEBUG] Removing corrupted metadata: {metadata_path}")
                os.remove(metadata_path)

    # If loading failed or fingerprints don't match, rebuild index
    print("[INFO] Building/rebuilding index from documents...")
    chunks = process_document_folder(
        data_folder,
        min_words_per_page=100,
        target_chunk_size=512,
        min_chunk_size=256,
        overlap_size=150
    )
    print(f"[DEBUG] Generated {len(chunks)} chunks")

    if not chunks:
        try:
            vectorstore = LC_FAISS.from_documents([Document(page_content="Empty index", metadata={})], embeddings)
            vectorstore.save_local(indexes_folder)
            with open(fingerprint_file, 'w') as f:
                json.dump(fingerprint_hash, f)
            print("[WARNING] No chunks generated, created empty index")
        except Exception as e:
            print(f"[ERROR] Failed to create empty vectorstore: {e}")
            raise
        return vectorstore

    try:
       # docs = [Document(page_content=ch["text"], metadata=ch["metadata"]) for ch in chunks]
        docs = [LCDocument(page_content=ch["text"], metadata=ch["metadata"]) for ch in chunks]
        vectorstore = LC_FAISS.from_documents(docs, embeddings)
        vectorstore.save_local(indexes_folder)
        with open(fingerprint_file, 'w') as f:
            json.dump(fingerprint_hash, f)
        print(f"[INFO] Created new vectorstore with {vectorstore.index.ntotal} vectors")
    except Exception as e:
        print(f"[ERROR] Failed to create vectorstore: {e}")
        raise
    return vectorstore


# ============================================================
# Шаблоны (prompt) для цепочек QA (отдельные для учителей и студентов)
# ============================================================
def get_teacher_prompt_template():
    return (
        "Ты — умный университетский ассистент для преподавателей и для сотрудников Университета AITU. Тебя зовут AITU - Connect."
        "Используй следующий полученный контекст для ответа на вопрос. "
        "Сначала разберись с заданным вопросом, попытайся его полностью понять, разберись с контекстом. Подумай перед тем как ответить"
        "Отвечай на том языке, на котором запрос."
        "Отвечай емко, четко и полно, в несколько абзацев, разбивая ответ на логически структурированные абзацы, "
        "Желательно, используй нумерованные списки или bullet points для наглядности. "
        "Будь дружелюбным, действуй как надежный друг и ассистент, старайся найти ответ на любой вопрос. Даже очень сложный. "
        "Если ты не знаешь ответа, честно скажи что не знаешь, но готов помочь с чем нибудь другим. "
        "Если в базе на вопрос нет данных или ответа, то предложи обратиться в Деканат (Dean’s Office) по электронной почте office.reg@astanait.edu.kz.\n\n"
        "Отвечай на том языке, на котором запрос. Если человек пишет на русском, то ответ должен быть на русском. Если человек пишет на английском, то и ответ должен быть на английском соответственно."
        "Если в базе данных недостаточно информации, сообщяй, что можно обратиться в офис регистратора.\n\n"
        "Предыдущее сообщение, обязательно держи это в памяти. Помни все, что обсуждали и спрашивали (история чата):\n"
        "{chat_history}\n\n"
        "Контекст документов:\n{context}\n\n"
        "Вопрос: {question}\n\n"
        "Ответ:"
    )

def get_student_prompt_template():
    return (
        "Ты — умный университетский ассистент для студентов Университета AITU. "
        "Тебя зовут AITU - Connect. "
        "Отвечай на том языке, на котором запрос. "
        "Сначала внимательно изучи вопрос, разберись в его сути и контексте. Подумай перед тем как ответить"
        "Отвечай емко, четко и полно, в несколько абзацев, разбивая ответ на логически структурированные абзацы, "
        "Желательно, используй нумерованные списки или bullet points для наглядности. "
        "Будь дружелюбным, отзывчивым и всегда старайся помочь студентам с их университетскими запросами. "
        "Если ты не знаешь ответа, честно скажи что не знаешь, но готов помочь с чем нибудь другим. "
        "Если в базе на вопрос нет данных или ответа, то предложи обратиться в Деканат (Dean’s Office) по электронной почте office.reg@astanait.edu.kz.\n\n"
        "Отвечай на том языке, на котором запрос. Если человек пишет на русском, то ответ должен быть на русском. Если человек пишет на английском, то и ответ должен быть на английском соответственно."
        "Предыдущее сообщение, обязательно держи это в памяти. Помни все, что обсуждали и спрашивали (история чата):\n"
        "{chat_history}\n\n"
        "Контекст:\n{context}\n\n"
        "Вопрос: {question}\n\n"
        "Ответ:"
    )

def get_teacher_flowchart_prompt():
    return (
        "Ты — умный университетский ассистент Университета AITU. "
        "На основании приведённого контекста и вопроса составь детальную и точную блок‑схему в формате Mermaid. "
        "Убедись, что итоговый синтаксис абсолютно корректен и готов к визуализации без правок.\n\n"
        "Контекст:\n{context}\n\n"
        "Вопрос: {question}\n\n"
        "Требования:\n"
        "- Выводи исключительно код Mermaid, начиная с ключевого слова `flowchart`.\n"
        "- Не добавляй никакого описательного текста и не оборачивай код в разметку (```), только чистый Mermaid.\n"
        "- Используй направление `TD` (сверху вниз) или `LR` (слева направо) в зависимости от структуры алгоритма.\n"
        "- Обозначь начало узлом `[Начало]`, конец — узлом `[Конец]`.\n"
        "- Для ветвлений, то есть условии, применяй ромбовидные узлы, фигура - ромб.\n"
        "- Для всех стрелок указывай текст условия или действия.\n\n"
        "Ответ (ТОЛЬКО Mermaid):"
    )


def get_student_flowchart_prompt():
    return (
        "Ты — умный университетский ассистент для студентов Университета AITU. "
        "На основании приведённого контекста и вопроса составь понятную и точную блок‑схему в формате Mermaid. "
        "Убедись, что итоговый синтаксис абсолютно корректен, без обёрток и лишнего текста.\n\n"
        "Контекст:\n{context}\n\n"
        "Вопрос: {question}\n\n"
        "Требования:\n"
        "- Выводи исключительно код Mermaid, начиная с ключевого слова `flowchart`.\n"
        "- Не добавляй никакого описательного текста и не оборачивай код в разметку (```), только чистый Mermaid.\n"
        "- Используй направление `TD` (сверху вниз) или `LR` (слева направо) в зависимости от структуры алгоритма.\n"
        "- Обозначь начало узлом `[Начало]`, конец — узлом `[Конец]`.\n"
        "- Для ветвлений, то есть условии, применяй ромбовидные узлы, фигура - ромб.\n"
        "- Для всех стрелок указывай текст условия или действия.\n\n"
        "Ответ (ТОЛЬКО Mermaid):"
    )

# ============================================================
# Инициализация LLM и цепочек QA для учителей и студентов
# ============================================================
llm = ChatOpenAI(openai_api_key=OPENAI_API_KEY, temperature=0, model_name="gpt-4o-mini")
teacher_prompt = PromptTemplate(
    template=get_teacher_prompt_template(),
    input_variables=["chat_history", "context", "question"]
)
student_prompt = PromptTemplate(
    template=get_student_prompt_template(),
    input_variables=["chat_history", "context", "question"]
)

teacher_vectorstore = load_or_rebuild_vectorstore(DATA_FOLDER_TEACHERS, INDEXES_FOLDER_TEACHERS)
student_vectorstore = load_or_rebuild_vectorstore(DATA_FOLDER_STUDENTS, INDEXES_FOLDER_STUDENTS)

teacher_qa_chain = ConversationalRetrievalChain.from_llm(
    llm,
    retriever=teacher_vectorstore.as_retriever(search_kwargs={"k": 3}),
    return_source_documents=True,
    combine_docs_chain_kwargs={"prompt": teacher_prompt}
)

student_qa_chain = ConversationalRetrievalChain.from_llm(
    llm,
    retriever=student_vectorstore.as_retriever(search_kwargs={"k": 3}),
    return_source_documents=True,
    combine_docs_chain_kwargs={"prompt": student_prompt}
)

# ============================================================
# Класс для управления историей чата с добавлением источников в ответ
# ============================================================
class ChatAssistant:
    def __init__(self, qa_chain):
        self.qa = qa_chain
        self.histories = {}  # {session_id: list of dicts}

    def get_answer(self, user_query: str, session_id: str = "default"):
        if session_id not in self.histories:
            self.histories[session_id] = []

        # 1) rewrite with history for follow‑ups
        chain_history = self._convert_history(session_id)

        # 2) call the chain
        result = self.qa({
            "question": user_query,
            "chat_history": chain_history
        })
        answer = result.get("answer", "")
        source_docs = result.get("source_documents", [])
        sources = self._extract_sources(source_docs)
        # if source_docs:
        #     sources = self._extract_sources(source_docs)
        #     if sources:
        #         sources_text = "\n\nSources:\n" + "\n".join(sources)
        #         answer += sources_text
     
        #self.histories[session_id].append(("user", user_query))
        #self.histories[session_id].append(("assistant", answer))
        # 3) record into history with timestamp

        self.histories[session_id].append({
            "role": "user",
            "content": user_query,
            "time": datetime.now().strftime("%I:%M %p"),
        })
        self.histories[session_id].append({
            "role": "assistant",
            "content": answer,
            "time": datetime.now().strftime("%I:%M %p"),
            "sources": sources
        })

        # 4) return pure answer + docs (your endpoint will convert docs→sources list)
        return answer, source_docs

    def _convert_history(self, session_id: str):
        history = self.histories.get(session_id, [])
        pairs = []
        last_user = None
        for entry in history:
            role = entry.get("role")
            content = entry.get("content")
            if role == "user":
                last_user = content
            elif role == "assistant" and last_user is not None:
                pairs.append((last_user, content))
                last_user = None
        return pairs
    def _extract_sources(self, source_docs) -> List[str]:
        seen = set()
        sources = []
        for doc in source_docs:
            file_name = doc.metadata.get("file_name")
            if file_name and file_name not in seen:
                seen.add(file_name)
                sources.append(file_name)
        return sources
    def clear_history(self, session_id: str = "default"):
        self.histories[session_id] = []

# Создаём объекты ChatAssistant для каждой модели
teacher_assistant = ChatAssistant(teacher_qa_chain)
student_assistant = ChatAssistant(student_qa_chain)

# ============================================================
# Утилита для извлечения списка источников (если нужно отдельно)
# ============================================================
def extract_sources_list(source_docs) -> List[str]:
    seen = set()
    sources = []
    for doc in source_docs:
        file_name = doc.metadata.get("file_name")
        if file_name and file_name not in seen:
            seen.add(file_name)
            sources.append(file_name)
    return sources

# ============================================================
# Инициализация менеджеров документов для учителей и студентов
# ============================================================
teacher_doc_manager = DocumentManager(DATA_FOLDER_TEACHERS)
student_doc_manager = DocumentManager(DATA_FOLDER_STUDENTS)

# ============================================================
# Модели данных для FastAPI
# ============================================================
class ChatRequest(BaseModel):
    query: str
    session_id: Optional[str] = "default"

class ChatResponse(BaseModel):
    answer: str
    sources: List[str] = []

# ============================================================
# Создание приложения FastAPI и эндпойнтов
# ============================================================
app = FastAPI(title="University Chat Assistant API")
from fastapi.staticfiles import StaticFiles

app.mount("/static", StaticFiles(directory="static", html=True), name="static")

# Эндпойнт для вывода всех маршрутов приложения
@app.get("/api/endpoints")
def list_endpoints():
    routes = []
    for route in app.routes:
        if hasattr(route, "methods"):
            methods = ", ".join(route.methods)
            routes.append({
                "path": route.path,
                "methods": methods,
                "name": route.name
            })
    return routes

# ---- Чат-эндпойнты для учителей и студентов ----
@app.post("/api/teacher/chat", response_model=ChatResponse)
def teacher_chat(payload: ChatRequest):
    answer, source_docs = teacher_assistant.get_answer(payload.query, payload.session_id)
    sources_list = extract_sources_list(source_docs)
    return ChatResponse(answer=answer, sources=sources_list)

@app.post("/api/student/chat", response_model=ChatResponse)
def student_chat(payload: ChatRequest):
    answer, source_docs = student_assistant.get_answer(payload.query, payload.session_id)
    sources_list = extract_sources_list(source_docs)
    return ChatResponse(answer=answer, sources=sources_list)

import re

    

@app.post("/api/teacher/flowchart")
def teacher_flowchart(payload: ChatRequest):
    # context
    relevant_docs = teacher_vectorstore.similarity_search(payload.query, k=3)
    context = "\n".join(doc.page_content for doc in relevant_docs)

    # sources 
    sources = extract_sources_list(relevant_docs)

    # LLM prompt → Mermaid code
    prompt = PromptTemplate(
        template=get_teacher_flowchart_prompt(),
        input_variables=["context", "question"]
    )
    chain = prompt | llm
    chain_response = chain.invoke({"context": context, "question": payload.query})
    mermaid_code = chain_response.content.strip()

    # debug
    print("[DEBUG] Mermaid code:\n", mermaid_code)
    print("[DEBUG] Sources:", sources)

    # JSON
    return JSONResponse({
        "mermaid": mermaid_code,
        "sources": sources
    })




@app.post("/api/student/flowchart")
def student_flowchart(payload: ChatRequest):
    # context vectorstore
    relevant_docs = student_vectorstore.similarity_search(payload.query, k=3)
    context = "\n".join(doc.page_content for doc in relevant_docs)

    # sources
    sources = extract_sources_list(relevant_docs)

    # student prompt
    prompt = PromptTemplate(
        template=get_student_flowchart_prompt(),
        input_variables=["context", "question"]
    )
    chain = prompt | llm
    chain_response = chain.invoke({"context": context, "question": payload.query})
    mermaid_code = chain_response.content.strip()

    # debug
    print("[DEBUG][STUDENT] Mermaid code:\n", mermaid_code)
    print("[DEBUG][STUDENT] Sources:", sources)

    # JSON envelope
    return JSONResponse({
        "mermaid": mermaid_code,
        "sources": sources
    })




# история чата

@app.get("/api/{role}/chat/clear")
def clear_chat(role: str, session_id: str = "default"):
    if role.lower() == "teacher":
        teacher_assistant.clear_history(session_id)
    elif role.lower() == "student":
        student_assistant.clear_history(session_id)
    else:
        raise HTTPException(status_code=404, detail="Role not found")
    return {"message": "История чата очищена"}

@app.get("/api/{role}/chat/history", response_model=ChatHistoryResponse)
def get_chat_history(role: str, session_id: str = "default"):
    if role.lower() == "teacher":
        hist = teacher_assistant.histories.get(session_id, [])
    elif role.lower() == "student":
        hist = student_assistant.histories.get(session_id, [])
    else:
        raise HTTPException(status_code=404, detail="Role not found")

    # timestamps
    conversation = [
        {
            "id": idx,
            "role": entry["role"],
            "content": entry["content"],
            "time": entry.get("time"),
            "sources": entry.get("sources", [])
        }
        for idx, entry in enumerate(hist)
    ]
    return {"session_id": session_id, "history": conversation}

@app.delete("/api/{role}/chat/history", response_model=ChatDeleteResponse)
def delete_chat_message(
    role: str,
    session_id: str = "default",
    message_id: int = None
):
    if role.lower() == "teacher":
        hist = teacher_assistant.histories.get(session_id, [])
    elif role.lower() == "student":
        hist = student_assistant.histories.get(session_id, [])
    else:
        raise HTTPException(status_code=404, detail="Role not found")

    if message_id is None or message_id < 0 or message_id >= len(hist):
        raise HTTPException(status_code=400, detail="Invalid message_id")

    hist.pop(message_id)
    return {"message": f"Deleted message {message_id} from session {session_id}"}




# ---- Эндпойнты для работы с документами (список и загрузка) ----
@app.get("/api/teacher/docs")
def list_teacher_docs():
    return teacher_doc_manager.get_active_documents()

@app.get("/api/student/docs")
def list_student_docs():
    return student_doc_manager.get_active_documents()



@app.post("/refresh/staff")
async def refresh_staff_index():
    try:
        global teacher_vectorstore, teacher_qa_chain
        print(f"[DEBUG] Starting refresh for staff index")
        teacher_vectorstore = load_or_rebuild_vectorstore(DATA_FOLDER_TEACHERS, INDEXES_FOLDER_TEACHERS,
                                                          call_id="refresh_staff")

        # Update the QA chain with the new retriever
        teacher_qa_chain.retriever = teacher_vectorstore.as_retriever(search_kwargs={"k": 3})

        print(f"[DEBUG] Rebuilt teacher_vectorstore with {teacher_vectorstore.index.ntotal} vectors")
        return {"message": "Индекс для сотрудников (Teacher) был успешно пересобран"}
    except Exception as e:
        print(f"[ERROR] Failed to refresh staff index: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/refresh/students")
def refresh_students_index():
    """
    Принудительно пересобрать индекс для студентов.
    """
    global student_vectorstore, student_qa_chain
    student_vectorstore = load_or_rebuild_vectorstore(DATA_FOLDER_STUDENTS, INDEXES_FOLDER_STUDENTS)

    # Update the QA chain with the new retriever
    student_qa_chain.retriever = student_vectorstore.as_retriever(search_kwargs={"k": 3})

    return {"message": "Индекс для студентов был успешно пересобран"}





@app.post("/api/{role}/docs/upload")
async def upload_doc(
    role: str,
    file: UploadFile = File(...),
    replace_doc_id: Optional[str] = Form(None)
):
    data_folder = DATA_FOLDER_TEACHERS if role=="teacher" else DATA_FOLDER_STUDENTS
    mgr = teacher_doc_manager if role=="teacher" else student_doc_manager

    content = await file.read()
    temp_path = os.path.join(data_folder, file.filename)
    with open(temp_path, "wb") as f:
        f.write(content)

    if replace_doc_id:
        mgr.delete_document_by_id(replace_doc_id)
        new_doc = mgr.add_document(temp_path)
        action = f"Replaced {replace_doc_id} → {new_doc['id']}"
    else:
        new_doc = mgr.add_document(temp_path)
        action = f"Added new document {new_doc['id']}"

    # rebuild index 
    idx_folder = INDEXES_FOLDER_TEACHERS if role=="teacher" else INDEXES_FOLDER_STUDENTS
    vs = load_or_rebuild_vectorstore(data_folder, idx_folder, call_id=f"upload_{role}")
    if role == "teacher":
        global teacher_vectorstore, teacher_qa_chain
        teacher_vectorstore = vs
        teacher_qa_chain.retriever = vs.as_retriever(search_kwargs={"k": 3})
    else:
        global student_vectorstore, student_qa_chain
        student_vectorstore = vs
        student_qa_chain.retriever = vs.as_retriever(search_kwargs={"k": 3})

    return {"message":"Done", "file_action":action}



# =====================
# Сначала с загруженного файла извлекает данные
# =====================
from docx import Document as DocxDocument

def extract_text_from_file(filepath):
    ext = os.path.splitext(filepath)[1].lower()
    if ext == '.txt':
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    elif ext == '.docx':
        doc = DocxDocument(filepath)
        return '\n'.join([p.text for p in doc.paragraphs])
    elif ext == '.pdf':
        text = ''
        with open(filepath, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                text += page.extract_text() or ''
        return text
    else:
        return ''

# =====================
# Cosine similarity
# =====================
def cosine_similarity(a, b):
    if not a.any() or not b.any():
        return 0.0
    return float(np.dot(a, b) / (np.linalg.norm(a) * np.linalg.norm(b)))

# =====================
# нормализация текста
# =====================
def normalize_text(text):
    import re
    text = text.lower().strip()
    text = re.sub(r'\s+', ' ', text)
    return text

# =====================
# схожести ищет
# =====================
def find_similar_files(uploaded_text, folder, threshold=0.7):
    model = SentenceTransformer('all-MiniLM-L12-v2')
    uploaded_text_norm = normalize_text(uploaded_text)
    uploaded_emb = model.encode([uploaded_text_norm], convert_to_numpy=True)[0]
    similar = []
    # print("[DEBUG] Uploaded text (first 200 chars):", uploaded_text_norm[:200])
    for fname in os.listdir(folder):
        fpath = os.path.join(folder, fname)
        if not os.path.isfile(fpath) or not fname.lower().endswith((".docx", ".pdf", ".txt")):
            continue
        try:
            text = extract_text_from_file(fpath)
            text_norm = normalize_text(text)
            if not text_norm.strip():
                continue
            emb = model.encode([text_norm], convert_to_numpy=True)[0]
            sim = cosine_similarity(uploaded_emb, emb)
            # print(f"[DEBUG] File: {fname} | Similarity: {sim:.4f} | Text (first 200 chars): {text_norm[:200]}")
            if sim >= threshold:
                similar.append({
                    'file': fname,
                    'similarity': round(sim * 100, 2)
                })
        except Exception as e:
            print(f"[ERROR] Processing {fname}: {e}")
            continue
    similar.sort(key=lambda x: x['similarity'], reverse=True)
    # если есть 100% similarity, то только его и показываем
    exact_matches = [f for f in similar if f['similarity'] == 100.0]
    if exact_matches:
        return exact_matches
    # ну либо топ 3
    return similar[:3]

@app.post("/api/{role}/docs/check_similarity")
async def check_doc_similarity(
    role: str,
    file: UploadFile = File(...)
):
    folder = DATA_FOLDER_TEACHERS if role=="teacher" else DATA_FOLDER_STUDENTS
    tmp = f"temp/{file.filename}"
    os.makedirs("temp", exist_ok=True)
    with open(tmp, "wb") as f: f.write(await file.read())

    text = extract_text_from_file(tmp)
    dups = find_similar_files(text, folder, threshold=0.7)
    os.remove(tmp)
    return {"possible_duplicates": dups}




if __name__ == "__main__":
    import uvicorn
    uvicorn.run("main:app", host="0.0.0.0", port=8000)
  # uvicorn.run("main:app", host="0.0.0.0", port=8000, reload=True)
